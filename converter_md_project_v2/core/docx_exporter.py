"""
Exportação de Markdown para DOCX.
"""

import io
import re

from docx import Document
from docx.shared import Inches


def _strip_frontmatter(text: str) -> str:
    """Remove bloco YAML frontmatter do início do texto."""
    lines = text.split("\n")
    if not lines or lines[0].strip() != "---":
        return text
    i = 1
    while i < len(lines):
        if lines[i].strip() == "---":
            return "\n".join(lines[i + 1:])
        i += 1
    return text


def _parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    """Extrai tabela markdown a partir da posição start.

    Returns:
        (rows, next_index) onde rows é lista de listas de células.
    """
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip()
        cells = [c.strip() for c in raw.strip("|").split("|")]
        if not all(re.match(r"^-+$", c) for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def markdown_to_docx(md_text: str, title: str = "") -> bytes:
    """Converte Markdown para DOCX.

    Args:
        md_text: Texto Markdown (pode conter frontmatter YAML).
        title: Título opcional para propriedades do documento.

    Returns:
        Bytes do arquivo DOCX.
    """
    clean_md = _strip_frontmatter(md_text)
    doc = Document()

    if title:
        doc.core_properties.title = title

    lines = clean_md.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Headings
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
            i += 1
            continue

        # Tabelas
        if stripped.startswith("|"):
            rows, i = _parse_table(lines, i)
            if rows:
                _add_table(doc, rows)
            continue

        # Blockquotes
        if stripped.startswith("> "):
            text = stripped[2:]
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(text)
            run.italic = True
            i += 1
            continue

        # Listas bullet
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            doc.add_paragraph(text, style="List Bullet")
            i += 1
            continue

        # Separador
        if stripped in ("---", "***"):
            doc.add_paragraph("─" * 50)
            i += 1
            continue

        # Parágrafo normal
        plain = _strip_inline_md(stripped)
        doc.add_paragraph(plain)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    """Adiciona tabela markdown ao documento DOCX."""
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < cols:
                table.rows[r_idx].cells[c_idx].text = cell_text


def _strip_inline_md(text: str) -> str:
    """Remove formatação inline markdown (bold, italic)."""
    text = re.sub(r"\*\*\*(.+?)\*\*\*", r"\1", text)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    return text
