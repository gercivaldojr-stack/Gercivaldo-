"""
Módulo de extração de texto de diferentes formatos de documentos.
Suporta: PDF, DOCX, TXT, Markdown.

Inclui:
- Extração nativa rápida via PyMuPDF
- OCR seletivo por página (apenas páginas com pouco texto nativo)
- Processamento em chunks para PDFs grandes (economia de RAM)
- Remoção de cabeçalhos/rodapés por análise de bbox
"""

import logging
import os
import re
import tempfile
from pathlib import Path

import chardet

logger = logging.getLogger(__name__)

# Padrões que identificam rodapés/cabeçalhos de escritório
_FOOTER_PATTERNS = [
    re.compile(r"CEP\s*\d{5}-?\d{3}", re.IGNORECASE),
    re.compile(r"\(\d{2}\)\s*\d[\s.\-]*\d{3,4}[\s.\-]*\d{4}"),
    re.compile(r"Página\s*\d+", re.IGNORECASE),
    re.compile(r"Pág\.\s*\d+", re.IGNORECASE),
    re.compile(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+"),
]


def extract_text(
    file_path: str = None,
    file_bytes: bytes | None = None,
    filename: str | None = None,
    preserve_inline_formatting: bool = True,
    ocr_enabled: bool = False,
    ocr_lang: str = "por",
    ocr_threshold: int = 30,
    page_range: str | None = None,
    chunk_size: int | None = None,
    detect_columns: bool = True,
    max_workers: int | None = None,
    ocr_cache_enabled: bool = False,
    ocr_cache_dir: str | None = None,
) -> str:
    """Extrai texto de um arquivo com base na extensão.

    Args:
        file_path: Caminho do arquivo no disco.
        file_bytes: Conteúdo do arquivo em bytes.
        filename: Nome original do arquivo.
        preserve_inline_formatting: Se True, preserva bold/italic do DOCX como Markdown.
        ocr_enabled: Se True, aplica OCR seletivo em páginas com pouco texto nativo.
        ocr_lang: Idioma do Tesseract para OCR (padrão: por = português).
        ocr_threshold: Mínimo de chars para considerar página com texto suficiente.
        page_range: Páginas a processar, 1-based (ex: "1-50", "1,5,10-20").
            Página 1 = primeira página. Apenas para PDF.
        chunk_size: Páginas por chunk para PDFs grandes. Processa N páginas por
            vez e libera memória entre chunks. None = processar tudo de uma vez.

    Returns:
        Texto extraído como string.
    """
    if file_bytes is not None and filename:
        ext = Path(filename).suffix.lower()
    elif file_path:
        ext = Path(file_path).suffix.lower()
    else:
        raise ValueError("Forneça file_path ou (file_bytes + filename)")

    extractors = {
        ".pdf": _extract_pdf,
        ".docx": _extract_docx,
        ".txt": _extract_txt,
        ".md": _extract_md,
    }

    extractor = extractors.get(ext)
    if not extractor:
        raise ValueError(f"Formato não suportado: {ext}. Use: {', '.join(extractors.keys())}")

    try:
        kwargs = {"preserve_inline_formatting": preserve_inline_formatting}
        if ext == ".pdf":
            kwargs.update(
                ocr_enabled=ocr_enabled,
                ocr_lang=ocr_lang,
                ocr_threshold=ocr_threshold,
                page_range=page_range,
                chunk_size=chunk_size,
                detect_columns=detect_columns,
                max_workers=max_workers,
                ocr_cache_enabled=ocr_cache_enabled,
                ocr_cache_dir=ocr_cache_dir,
            )

        if file_bytes is not None:
            return extractor(file_bytes=file_bytes, file_path=None, **kwargs)
        else:
            with open(file_path, "rb") as f:
                raw = f.read()
            return extractor(file_bytes=raw, file_path=file_path, **kwargs)
    except Exception as e:
        logger.error("Erro ao extrair texto de %s: %s", filename or file_path, e)
        raise


# ============================================================
# PDF extraction with table support and footer removal via bbox
# ============================================================

def _is_footer_text(text: str) -> bool:
    """Verifica se um texto corresponde a padrões de rodapé."""
    return any(p.search(text) for p in _FOOTER_PATTERNS)


def _detect_hf_zones(doc, sample_pages: int = 50) -> set[tuple[int, int]]:
    """Detecta zonas de cabeçalho/rodapé por AMOSTRAGEM.

    Em vez de escanear TODAS as páginas, faz amostragem inteligente:
    * Se doc tem <= sample_pages páginas: escaneia todas
    * Se doc tem > sample_pages: amostra uniformemente sample_pages páginas
      (primeiras 10 + últimas 10 + N uniformemente do meio)
    * Usa threshold proporcional à amostra (50% das páginas amostradas)
    * Após identificar padrões na amostra, aplica a todas as páginas
      de forma streaming (uma página por vez, liberando memória)

    Isso reduz memória de O(N*page_size) para O(sample*page_size).

    Retorna set de (page_index, block_index) a remover.
    """
    total_pages = len(doc)
    if total_pages < 3:
        return set()

    # Determinar páginas a amostrar
    if total_pages <= sample_pages:
        sampled_indices = list(range(total_pages))
    else:
        head = list(range(min(10, total_pages)))
        tail = list(range(max(total_pages - 10, 0), total_pages))
        middle_count = sample_pages - len(head) - len(tail)
        if middle_count > 0:
            step = (total_pages - 20) / (middle_count + 1)
            middle = [int(10 + step * (i + 1)) for i in range(middle_count)]
        else:
            middle = []
        sampled_indices = sorted(set(head + middle + tail))
        logger.debug(
            "hf_zones: amostrando %d de %d páginas",
            len(sampled_indices), total_pages,
        )

    # Fase 1: Coletar y-positions APENAS das páginas amostradas
    y_positions: dict[int, list[tuple[int, int, str]]] = {}

    for page_idx in sampled_indices:
        page = doc[page_idx]
        page_height = page.rect.height
        blocks = page.get_text("dict")["blocks"]
        for blk_idx, block in enumerate(blocks):
            if block["type"] != 0:
                continue
            bbox = block["bbox"]
            y_top = bbox[1]
            y_bottom = bbox[3]

            in_header_zone = y_top < page_height * 0.12
            in_footer_zone = y_bottom > page_height * 0.88
            if not in_header_zone and not in_footer_zone:
                continue

            block_text = ""
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    block_text += span.get("text", "")
            block_text = block_text.strip()
            if not block_text:
                continue

            y_key = round(y_top / 5) * 5
            y_positions.setdefault(y_key, []).append(
                (page_idx, blk_idx, block_text)
            )
        del blocks

    # Fase 2: Identificar padrões com threshold proporcional à amostra
    threshold = len(sampled_indices) * 0.5
    final_remove_set: set[tuple[int, int]] = set()
    identified_y_keys: set[int] = set()

    for y_key, entries in y_positions.items():
        pages_with_this_y = {e[0] for e in entries}
        if len(pages_with_this_y) >= threshold:
            identified_y_keys.add(y_key)
            for page_idx, blk_idx, text in entries:
                final_remove_set.add((page_idx, blk_idx))
                logger.debug(
                    "Removendo bloco hf: página %d, y=%d, text='%s'",
                    page_idx, y_key, text[:50],
                )

    # Fase 3: Aplicar padrões a páginas NÃO amostradas (streaming)
    sampled_set = set(sampled_indices)
    if identified_y_keys:
        for page_idx in range(total_pages):
            if page_idx in sampled_set:
                continue
            page = doc[page_idx]
            page_height = page.rect.height
            blocks = page.get_text("dict")["blocks"]
            for blk_idx, block in enumerate(blocks):
                if block["type"] != 0:
                    continue
                bbox = block["bbox"]
                y_top = bbox[1]
                y_bottom = bbox[3]

                in_header_zone = y_top < page_height * 0.12
                in_footer_zone = y_bottom > page_height * 0.88
                if not in_header_zone and not in_footer_zone:
                    continue

                y_key = round(y_top / 5) * 5
                if y_key in identified_y_keys:
                    final_remove_set.add((page_idx, blk_idx))

                # Também verificar padrões de rodapé por regex
                block_text = ""
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        block_text += span.get("text", "")
                block_text = block_text.strip()
                if (
                    block_text
                    and _is_footer_text(block_text)
                    and y_bottom > page_height * 0.85
                ):
                    final_remove_set.add((page_idx, blk_idx))
            del blocks

    # Footer regex para páginas amostradas também
    for page_idx in sampled_indices:
        page = doc[page_idx]
        page_height = page.rect.height
        blocks = page.get_text("dict")["blocks"]
        for blk_idx, block in enumerate(blocks):
            if block["type"] != 0:
                continue
            bbox = block["bbox"]
            y_bottom = bbox[3]
            if y_bottom <= page_height * 0.85:
                continue
            block_text = ""
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    block_text += span.get("text", "")
            block_text = block_text.strip()
            if block_text and _is_footer_text(block_text):
                final_remove_set.add((page_idx, blk_idx))
        del blocks

    return final_remove_set


def _table_to_markdown(table) -> str:
    """Converte uma tabela PyMuPDF para formato Markdown."""
    try:
        data = table.extract()
    except Exception:
        return ""

    if not data or not data[0]:
        return ""

    lines = []
    # Header
    header = [str(cell or "").strip().replace("\n", " ") for cell in data[0]]
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join("---" for _ in header) + " |")
    # Rows
    for row in data[1:]:
        cells = [str(cell or "").strip().replace("\n", " ") for cell in row]
        # Pad se necessário
        while len(cells) < len(header):
            cells.append("")
        lines.append("| " + " | ".join(cells[:len(header)]) + " |")

    return "\n".join(lines)


def _parse_page_range(spec: str, total_pages: int) -> list[int]:
    """Converte especificação de páginas 1-based do usuário em lista 0-based.

    Entrada do usuário é 1-based (página 1 = primeira página).
    Saída é 0-based (índice 0 = primeira página) para uso interno com PyMuPDF.

    Exemplos (documento com 20 páginas):
        "1"       → [0]
        "1-10"    → [0, 1, 2, ..., 9]
        "1,5,10-20" → [0, 4, 9, 10, ..., 19]

    Páginas fora do limite são silenciosamente ignoradas.
    Entradas malformadas geram ValueError.

    Args:
        spec: String com páginas 1-based. Ex: "1", "1-10", "1,5,10-20".
        total_pages: Número total de páginas do documento.

    Returns:
        Lista ordenada de índices 0-based.

    Raises:
        ValueError: Se a spec contém partes não-numéricas ou vazias.
    """
    if not spec or not spec.strip():
        return list(range(total_pages))

    pages = set()
    for part in spec.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            pieces = part.split("-", 1)
            try:
                start_1 = int(pieces[0])
                end_1 = int(pieces[1])
            except ValueError:
                raise ValueError(f"Intervalo de paginas invalido: '{part}'. Use formato N-M (ex: 1-10).")
            if start_1 < 1:
                raise ValueError(f"Pagina deve ser >= 1, recebeu {start_1}.")
            if end_1 < start_1:
                raise ValueError(f"Intervalo invalido: {start_1}-{end_1} (inicio > fim).")
            # Converter 1-based para 0-based; limitar ao total
            start_0 = start_1 - 1
            end_0 = min(end_1 - 1, total_pages - 1)
            if start_0 < total_pages:
                pages.update(range(start_0, end_0 + 1))
        else:
            try:
                p_1 = int(part)
            except ValueError:
                raise ValueError(f"Pagina invalida: '{part}'. Use numero inteiro (ex: 5).")
            if p_1 < 1:
                raise ValueError(f"Pagina deve ser >= 1, recebeu {p_1}.")
            p_0 = p_1 - 1
            if p_0 < total_pages:
                pages.add(p_0)

    return sorted(pages)


def _ocr_page(page, lang: str = "por") -> str:
    """Aplica OCR em uma página PDF via pytesseract.

    Renderiza a página como imagem em resolução moderada (200 DPI)
    para equilibrar qualidade e uso de memória.
    """
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        logger.warning("pytesseract ou Pillow nao instalado. OCR desabilitado.")
        return ""

    try:
        # 200 DPI: bom equilíbrio entre qualidade OCR e uso de memória
        mat = page.get_pixmap(dpi=200)
        img = Image.frombytes("RGB", (mat.width, mat.height), mat.samples)
        text = pytesseract.image_to_string(img, lang=lang)
        # Liberar memória imediatamente
        del mat, img
        return text
    except Exception as e:
        logger.warning("OCR falhou na pagina: %s", e)
        return ""


def _get_ocr_cached(page, lang: str, ocr_cache) -> str:
    """OCR com cache: tenta cache primeiro, senão chama _ocr_page e salva."""
    if ocr_cache is not None:
        cached = ocr_cache.get(page, lang)
        if cached is not None:
            return cached
    text = _ocr_page(page, lang=lang)
    if ocr_cache is not None and text:
        ocr_cache.put(page, lang, text)
    return text


def _extract_single_page(page, page_idx: int, remove_set: set,
                         ocr_enabled: bool, ocr_lang: str,
                         ocr_threshold: int,
                         detect_columns: bool = True,
                         ocr_cache=None) -> tuple[str, bool]:
    """Extrai texto de uma única página PDF.

    Returns:
        Tupla (texto_da_pagina, usou_ocr).
    """
    # Fallback: se "auto" chegou aqui (ex: worker paralelo), resolver por página
    if ocr_lang == "auto":
        from .lang_detector import detect_language_from_page
        ocr_lang = detect_language_from_page(page)

    page_parts = []
    used_ocr = False

    # Detecção de colunas: se habilitado, usa ordem de leitura inteligente
    if detect_columns:
        from .column_detector import (
            _is_two_column_layout,
            detect_and_reorder_columns,
        )
        blocks_raw = page.get_text("dict")["blocks"]
        text_blocks = []
        for blk in blocks_raw:
            if blk["type"] != 0:
                continue
            txt = ""
            for ln in blk.get("lines", []):
                for sp in ln.get("spans", []):
                    txt += sp.get("text", "")
            if txt.strip():
                text_blocks.append({
                    "x0": blk["bbox"][0], "y0": blk["bbox"][1],
                    "x1": blk["bbox"][2], "y1": blk["bbox"][3],
                    "text": txt.strip(),
                })
        if _is_two_column_layout(text_blocks, page.rect.width):
            page_text = detect_and_reorder_columns(page)
            if ocr_enabled and len(page_text.strip()) < ocr_threshold:
                ocr_text = _get_ocr_cached(
                    page, ocr_lang, ocr_cache
                )
                if ocr_text and ocr_text.strip():
                    page_text = ocr_text
                    used_ocr = True
            return page_text, used_ocr

    # Tentar extrair tabelas
    tables = []
    table_rects = []
    try:
        found_tables = page.find_tables()
        if found_tables and found_tables.tables:
            for tbl in found_tables.tables:
                md_table = _table_to_markdown(tbl)
                if md_table:
                    tables.append(md_table)
                    table_rects.append(tbl.bbox)
    except Exception as e:
        logger.debug("find_tables() falhou na página %d: %s", page_idx + 1, e)

    # Extrair texto via dict para controle de bbox
    blocks = page.get_text("dict")["blocks"]
    for blk_idx, block in enumerate(blocks):
        if block["type"] != 0:
            continue

        if (page_idx, blk_idx) in remove_set:
            continue

        blk_bbox = block["bbox"]
        in_table = False
        for trect in table_rects:
            if (blk_bbox[1] >= trect[1] - 2 and blk_bbox[3] <= trect[3] + 2
                    and blk_bbox[0] >= trect[0] - 2 and blk_bbox[2] <= trect[2] + 2):
                in_table = True
                break
        if in_table:
            continue

        block_text = ""
        for line in block.get("lines", []):
            line_text = ""
            for span in line.get("spans", []):
                line_text += span.get("text", "")
            if line_text.strip():
                block_text += line_text + "\n"

        if block_text.strip():
            page_parts.append(block_text.strip())

    for md_table in tables:
        page_parts.append("\n" + md_table + "\n")

    page_text = "\n".join(page_parts)

    # OCR seletivo: só aplica se a página tem pouco texto nativo
    if ocr_enabled and len(page_text.strip()) < ocr_threshold:
        logger.debug("Página %d com %d chars (< %d): aplicando OCR",
                     page_idx + 1, len(page_text.strip()), ocr_threshold)
        ocr_text = _get_ocr_cached(page, ocr_lang, ocr_cache)
        if ocr_text and ocr_text.strip():
            page_text = ocr_text
            used_ocr = True

    return page_text, used_ocr


def _extract_pdf(
    file_bytes: bytes,
    file_path: str | None = None,
    ocr_enabled: bool = False,
    ocr_lang: str = "por",
    ocr_threshold: int = 30,
    page_range: str | None = None,
    chunk_size: int | None = None,
    detect_columns: bool = True,
    max_workers: int | None = None,
    ocr_cache_enabled: bool = False,
    ocr_cache_dir: str | None = None,
    **kwargs,
) -> str:
    """Extrai texto de PDF usando PyMuPDF (fitz).

    Funcionalidades:
    - Extração de tabelas via find_tables() com fallback para texto puro.
    - Remoção de cabeçalhos/rodapés via análise de bbox.
    - OCR seletivo: só aplica OCR em páginas com menos de ocr_threshold chars.
    - page_range: entrada 1-based do usuário (ex: "1-10") para processar
      apenas um subconjunto de páginas.
    - chunk_size: processa N páginas por vez, fechando e reabrindo o doc
      entre chunks para liberar memória. Útil para PDFs com 500+ páginas
      em máquinas com pouca RAM.
    """
    import gc

    import fitz

    text_parts = []
    tmp_path = None

    try:
        if file_path and os.path.exists(file_path):
            doc = fitz.open(file_path)
        else:
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
            tmp.write(file_bytes)
            tmp.close()
            tmp_path = tmp.name
            doc = fitz.open(tmp_path)

        total_pages = len(doc)
        logger.info("PDF com %d páginas", total_pages)

        # Determinar quais páginas processar (entrada 1-based, saída 0-based)
        if page_range:
            pages_to_process = _parse_page_range(page_range, total_pages)
            logger.info("Processando %d de %d páginas (range: %s)",
                        len(pages_to_process), total_pages, page_range)
        else:
            pages_to_process = list(range(total_pages))

        if not pages_to_process:
            logger.warning("Nenhuma página a processar (range fora do limite?)")
            doc.close()
            return ""

        # Detectar zonas de cabeçalho/rodapé (scan leve sobre todo o doc)
        remove_set = _detect_hf_zones(doc)
        if remove_set:
            logger.info("Detectados %d blocos de cabeçalho/rodapé para remoção",
                        len(remove_set))

        # Detectar idioma automaticamente se ocr_lang == "auto"
        if ocr_enabled and ocr_lang == "auto":
            from .lang_detector import detect_document_language
            ocr_lang = detect_document_language(doc)
            logger.info("Idioma detectado automaticamente: %s", ocr_lang)

        # Criar instância de OCRCache se habilitado
        ocr_cache = None
        if ocr_enabled and ocr_cache_enabled:
            from .ocr_cache import OCRCache
            ocr_cache = OCRCache(
                cache_dir=ocr_cache_dir, enabled=True
            )

        # Fechar doc antes do loop de chunks — será reaberto por chunk
        # Isso só faz sentido se chunk_size está definido; caso contrário
        # mantemos o doc aberto para o loop simples.
        doc_path = file_path if (file_path and os.path.exists(file_path)) else tmp_path
        doc.close()

        # ── Processamento paralelo de chunks (se disponível) ──
        if (
            chunk_size and chunk_size > 0
            and max_workers != 1
            and len(pages_to_process) > chunk_size
        ):
            from .parallel import process_pdf_chunks_parallel
            par_result, par_ocr = process_pdf_chunks_parallel(
                doc_path=doc_path,
                pages_to_process=pages_to_process,
                chunk_size=chunk_size,
                remove_set=remove_set,
                ocr_enabled=ocr_enabled,
                ocr_lang=ocr_lang,
                ocr_threshold=ocr_threshold,
                detect_columns=detect_columns,
                max_workers=max_workers,
                ocr_cache_enabled=ocr_cache_enabled,
                ocr_cache_dir=ocr_cache_dir,
            )
            if par_result is not None:
                if par_ocr > 0:
                    logger.info(
                        "OCR aplicado em %d páginas (paralelo)",
                        par_ocr,
                    )
                return "\n".join(par_result)

        # ── Processamento sequencial em chunks ──
        effective_chunk = (
            chunk_size
            if (chunk_size and chunk_size > 0)
            else len(pages_to_process)
        )
        ocr_count = 0
        chunks_processed = 0

        for chunk_start in range(0, len(pages_to_process), effective_chunk):
            chunk_pages = pages_to_process[chunk_start:chunk_start + effective_chunk]
            chunks_processed += 1

            if chunk_size and chunk_size > 0:
                logger.info("Chunk %d: páginas %s (%d págs)",
                            chunks_processed,
                            f"{chunk_pages[0]+1}-{chunk_pages[-1]+1}",
                            len(chunk_pages))

            # Abrir doc para este chunk
            chunk_doc = fitz.open(doc_path)

            for page_idx in chunk_pages:
                page = chunk_doc[page_idx]
                try:
                    page_text, used_ocr = _extract_single_page(
                        page, page_idx, remove_set,
                        ocr_enabled, ocr_lang, ocr_threshold,
                        detect_columns=detect_columns,
                        ocr_cache=ocr_cache,
                    )
                    if used_ocr:
                        ocr_count += 1
                    if page_text.strip():
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning("Erro na página %d: %s", page_idx + 1, e)
                    try:
                        fallback = page.get_text("text")
                        if fallback.strip():
                            text_parts.append(fallback)
                    except Exception:
                        text_parts.append(f"\n[Erro ao extrair página {page_idx + 1}]\n")

            # Fechar doc do chunk e forçar GC para liberar memória
            chunk_doc.close()
            if chunk_size and chunk_size > 0:
                gc.collect()

        if ocr_count > 0:
            logger.info("OCR aplicado em %d de %d páginas",
                        ocr_count, len(pages_to_process))
        if chunks_processed > 1:
            logger.info("Processamento concluído em %d chunks de %d páginas",
                        chunks_processed, chunk_size)

    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)

    return "\n".join(text_parts)


# ============================================================
# DOCX extraction with table support via iter_block_items
# ============================================================

def _iter_block_items(parent):
    """Itera sobre parágrafos e tabelas na ordem em que aparecem no documento DOCX."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn

    for child in parent.element.body:
        if child.tag == qn('w:p'):
            yield Paragraph(child, parent)
        elif child.tag == qn('w:tbl'):
            yield Table(child, parent)


def _docx_table_to_markdown(table) -> str:
    """Converte uma tabela python-docx para formato Markdown."""
    rows = table.rows
    if not rows:
        return ""

    lines = []
    # Header (primeira linha)
    header_cells = [cell.text.strip().replace("\n", " ") for cell in rows[0].cells]
    lines.append("| " + " | ".join(header_cells) + " |")
    lines.append("| " + " | ".join("---" for _ in header_cells) + " |")

    # Data rows
    for row in list(rows)[1:]:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        while len(cells) < len(header_cells):
            cells.append("")
        lines.append("| " + " | ".join(cells[:len(header_cells)]) + " |")

    return "\n".join(lines)


def _paragraph_to_markdown(paragraph, preserve_formatting: bool = True) -> str:
    """Converte um parágrafo DOCX para Markdown preservando bold/italic.

    Itera sobre runs do parágrafo e envolve trechos em **...** ou *...*.
    """
    if not preserve_formatting:
        return paragraph.text.strip()

    parts = []
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
        bold = run.bold
        italic = run.italic
        if bold and italic:
            parts.append(f"***{text}***")
        elif bold:
            parts.append(f"**{text}**")
        elif italic:
            parts.append(f"*{text}*")
        else:
            parts.append(text)

    result = "".join(parts).strip()
    # Limpar marcadores vazios: ****, **** etc.
    result = re.sub(r"\*{2,3}\s*\*{2,3}", "", result)
    # Mesclar marcadores adjacentes: **texto****outro** → **textooutro**
    result = re.sub(r"\*\*\*\*\*\*", "", result)  # ****** vazio
    result = re.sub(r"\*\*\*\*", "", result)        # **** entre bolds adjacentes
    result = re.sub(r"\*\*\s+\*\*", " ", result)   # ** ** espaço entre bolds
    return result


def _extract_docx(
    file_bytes: bytes,
    file_path: str | None = None,
    preserve_inline_formatting: bool = True,
    **kwargs,
) -> str:
    """Extrai texto de DOCX preservando estrutura de parágrafos, tabelas e formatação inline."""
    import io

    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = Document(io.BytesIO(file_bytes))
    parts = []

    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = _paragraph_to_markdown(block, preserve_inline_formatting)
            if not text:
                continue

            style_name = (block.style.name or "").lower() if block.style else ""

            if "heading 1" in style_name:
                # Strip inline formatting from headings (redundant with #)
                plain = block.text.strip()
                parts.append(f"# {plain}")
            elif "heading 2" in style_name:
                plain = block.text.strip()
                parts.append(f"## {plain}")
            elif "heading 3" in style_name:
                plain = block.text.strip()
                parts.append(f"### {plain}")
            elif "heading 4" in style_name:
                plain = block.text.strip()
                parts.append(f"#### {plain}")
            elif "title" in style_name:
                plain = block.text.strip()
                parts.append(f"# {plain}")
            else:
                parts.append(text)

        elif isinstance(block, Table):
            md_table = _docx_table_to_markdown(block)
            if md_table:
                parts.append("")
                parts.append(md_table)
                parts.append("")

    return "\n".join(parts)


def _extract_txt(file_bytes: bytes, file_path: str | None = None, **kwargs) -> str:
    """Extrai texto de arquivo TXT detectando encoding automaticamente."""
    detected = chardet.detect(file_bytes)
    encoding = detected.get("encoding", "utf-8") or "utf-8"

    try:
        return file_bytes.decode(encoding)
    except (UnicodeDecodeError, LookupError):
        logger.warning("Fallback de encoding: %s -> utf-8", encoding)
        return file_bytes.decode("utf-8", errors="replace")


def _extract_md(file_bytes: bytes, file_path: str | None = None, **kwargs) -> str:
    """Retorna conteúdo Markdown como está (já é o formato alvo)."""
    return _extract_txt(file_bytes, file_path)
