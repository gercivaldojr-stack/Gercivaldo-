"""Script de debug para pipeline DOCX real.

Constrói um DOCX simulando "Capítulo 11. Serviços Públicos" de Gabriel Pires,
roda o pipeline completo em modo doutrina e reporta PASS/FAIL para cada
problema P1-P10 com evidência concreta.
"""

import io
import sys

sys.path.insert(0, '.')

from docx import Document

from core.pipeline import convert_document


def build_real_docx() -> bytes:
    """Constrói DOCX realista simulando capítulo do JusBrasil."""
    doc = Document()

    # Bloco bibliográfico (primeira coisa no doc)
    doc.add_paragraph(
        "PIRES, Gabriel. Manual de Direito Administrativo. Brasília: JusBrasil, "
        "2023. Capítulo 11. Serviços Públicos. In: JusBrasil. Disponível em: "
        "httpswwwjaborandi.jusbrasil.com.br/cap-11. Acesso em: 15 out. 2023."
    )

    # Título H1
    doc.add_heading("Capítulo 11. Serviços Públicos", level=1)

    # UI text
    doc.add_paragraph("Clique aqui e acesse o vídeo sobre o tema.")

    # Heading H2 fundido com corpo (caso real)
    doc.add_heading(
        "11.1. Relevância e atualidade da questão A noção jurídica de serviço público",
        level=2,
    )

    # Parágrafo com nota de rodapé inline colada
    doc.add_paragraph(
        "O serviço público1 é toda atividade prestada pela Administração "
        "Pública com o objetivo de satisfazer necessidades essenciais da "
        "coletividade2, sob regime de direito público."
    )

    # Seção numerada como paragrafo (deve virar heading)
    doc.add_paragraph("11.2. Conceito")
    doc.add_paragraph("O serviço público é toda atividade prestada pela Administração.")

    # Enumerações coladas
    doc.add_paragraph("a)serviço comum")
    doc.add_paragraph("b)taxa administrativa")
    doc.add_paragraph("c)preço público")

    # Parágrafo longo (>500 chars com 4+ frases)
    doc.add_paragraph(
        "A prestação de serviço público é regulamentada por diversos "
        "dispositivos legais brasileiros. O Artigo 175 da Constituição "
        "Federal estabelece princípios fundamentais para esta atividade. "
        "A Lei 8.987 de 1995 regula a concessão e permissão de serviços. "
        "Já a Lei 11.079 de 2004 trata das parcerias público-privadas. "
        "Cada diploma normativo estabelece regras específicas para a "
        "modalidade abordada, respeitando sempre os princípios basilares "
        "da administração pública."
    )

    # URL malformada isolada
    doc.add_paragraph(
        "Mais informações em httpswwwjaborandi.jusbrasil.com.br/servicos"
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def check(name: str, passed: bool, evidence="") -> bool:
    """Imprime PASS/FAIL com evidência."""
    status = "PASS" if passed else "FAIL"
    print(f"  [{status}] {name}")
    if not passed and evidence:
        if isinstance(evidence, list):
            ev_str = "\n".join(str(x) for x in evidence)
        else:
            ev_str = str(evidence)
        for line in ev_str.split("\n")[:5]:
            print(f"         └─ {line[:100]}")
    return passed


def main():
    print("=" * 70)
    print("DEBUG: Pipeline DOCX real (Capítulo 11 - Serviços Públicos)")
    print("=" * 70)

    docx_bytes = build_real_docx()
    print(f"\nDOCX construído: {len(docx_bytes)} bytes\n")

    result = convert_document(
        file_bytes=docx_bytes,
        filename="Capitulo_11_Servicos_Publicos.docx",
        mode="doutrina",
    )

    if not result.success:
        print(f"ERRO na conversão: {result.error}")
        return 1

    md = result.markdown
    print("\n" + "=" * 70)
    print("SAÍDA DO PIPELINE:")
    print("=" * 70)
    print(md)
    print("=" * 70)

    # Separar frontmatter e body para checks
    lines = md.split("\n")
    in_fm = False
    body_lines = []
    fm_lines = []
    for ln in lines:
        if ln.strip() == "---":
            in_fm = not in_fm
            fm_lines.append(ln)
            continue
        if in_fm:
            fm_lines.append(ln)
        else:
            body_lines.append(ln)
    frontmatter = "\n".join(fm_lines)
    body = "\n".join(body_lines)

    print("\n" + "=" * 70)
    print("CHECKS P1-P10:")
    print("=" * 70)

    passed_count = 0
    total = 10

    # P1: título NÃO deve ser referência bibliográfica
    fm_has_biblio = (
        "In:" in frontmatter or "Acesso em" in frontmatter
        or 'titulo: "PIRES, Gabriel' in frontmatter
    )
    passed_count += check(
        "P1: Título não é referência bibliográfica",
        not fm_has_biblio,
        evidence=frontmatter,
    )

    # P2: corpo NÃO deve conter bloco bibliográfico
    body_has_biblio = (
        "PIRES, Gabriel" in body and "Acesso em" in body
    ) or "Acesso em: 15 out. 2023" in body
    passed_count += check(
        "P2: Corpo sem bloco bibliográfico",
        not body_has_biblio,
        evidence=body[:500],
    )

    # P3: corpo NÃO deve ter "CAPÍTULO 11" solto (TOC residual)
    # Verificar se "Capítulo 11" aparece SEM heading # (texto solto)
    toc_residual = False
    for ln in body_lines:
        s = ln.strip()
        if s.startswith("CAPÍTULO 11") and not s.startswith("#"):
            toc_residual = True
            break
    passed_count += check(
        "P3: Sem 'CAPÍTULO 11' solto (TOC residual)",
        not toc_residual,
    )

    # P4: headings NÃO devem estar fundidos com corpo
    fused_heading = False
    fused_line = ""
    for ln in body_lines:
        s = ln.strip()
        if s.startswith(("## ", "### ", "#### ")):
            # Se contém mais de 15 palavras, provável fusão
            content = s.split(" ", 1)[1] if " " in s else ""
            words = content.split()
            if len(words) > 15:
                fused_heading = True
                fused_line = s
                break
    passed_count += check(
        "P4: Headings não fundidos com corpo",
        not fused_heading,
        evidence=fused_line,
    )

    # P5: "11.2. Conceito" deve virar heading
    p5_found = any(
        ln.strip().startswith("#") and "11.2" in ln
        for ln in body_lines
    )
    passed_count += check(
        "P5: '11.2. Conceito' é heading",
        p5_found,
        evidence=body[:500],
    )

    # P6: "a)serviço" deve virar "a) serviço"
    p6_ok = (
        "a)serviço" not in body and "b)taxa" not in body
    )
    passed_count += check(
        "P6: Enumerações com espaço (a) servico)",
        p6_ok,
        evidence=body,
    )

    # P7: "httpswww" deve virar "https://www."
    p7_ok = "httpswww" not in body
    passed_count += check(
        "P7: URLs malformadas corrigidas",
        p7_ok,
        evidence=[ln for ln in body_lines if "jaborandi" in ln.lower()][:2],
    )

    # P8: "Clique aqui" deve ser removido
    p8_ok = "Clique aqui e acesse" not in body
    passed_count += check(
        "P8: Texto UI 'Clique aqui' removido",
        p8_ok,
        evidence=[ln for ln in body_lines if "Clique" in ln][:2],
    )

    # P9: parágrafos >400 chars devem ser divididos
    long_para = ""
    p9_ok = True
    for ln in body_lines:
        s = ln.strip()
        if (
            len(s) > 500
            and not s.startswith(("#", ">", "|", "-", "*"))
        ):
            p9_ok = False
            long_para = s
            break
    passed_count += check(
        "P9: Parágrafos >500 chars divididos",
        p9_ok,
        evidence=f"[{len(long_para)} chars] {long_para[:150]}",
    )

    # P10: "público1" deve virar "público[^1]" ou remover
    p10_ok = "público1" not in body
    passed_count += check(
        "P10: Notas de rodapé inline tratadas",
        p10_ok,
        evidence=[ln for ln in body_lines if "público1" in ln][:2],
    )

    print("\n" + "=" * 70)
    print(f"RESUMO: {passed_count}/{total} PASS")
    print("=" * 70)
    return 0 if passed_count == total else 1


if __name__ == "__main__":
    sys.exit(main())
