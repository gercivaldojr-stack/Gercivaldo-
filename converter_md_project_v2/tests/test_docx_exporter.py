"""Testes para exportação DOCX."""

import io

from docx import Document

from core.docx_exporter import markdown_to_docx


class TestMarkdownToDocx:
    def test_basic_conversion(self):
        md = "# Titulo\n\nParagrafo de texto."
        result = markdown_to_docx(md)
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_output_is_valid_docx(self):
        md = "# Test\n\nTexto."
        result = markdown_to_docx(md)
        doc = Document(io.BytesIO(result))
        assert len(doc.paragraphs) > 0

    def test_headings(self):
        md = "# H1\n\n## H2\n\n### H3\n\nTexto."
        result = markdown_to_docx(md)
        doc = Document(io.BytesIO(result))
        styles = [p.style.name for p in doc.paragraphs]
        assert "Heading 1" in styles
        assert "Heading 2" in styles
        assert "Heading 3" in styles

    def test_tables(self):
        md = "| Col1 | Col2 |\n| --- | --- |\n| A | B |"
        result = markdown_to_docx(md)
        doc = Document(io.BytesIO(result))
        assert len(doc.tables) == 1
        assert doc.tables[0].rows[0].cells[0].text == "Col1"

    def test_blockquotes(self):
        md = "> Citacao juridica."
        result = markdown_to_docx(md)
        doc = Document(io.BytesIO(result))
        found = any("Citacao juridica" in p.text for p in doc.paragraphs)
        assert found

    def test_title_in_properties(self):
        md = "# Doc\n\nTexto."
        result = markdown_to_docx(md, title="Meu Titulo")
        doc = Document(io.BytesIO(result))
        assert doc.core_properties.title == "Meu Titulo"

    def test_frontmatter_removed(self):
        md = "---\ntitulo: Test\n---\n\n# Doc\n\nTexto."
        result = markdown_to_docx(md)
        doc = Document(io.BytesIO(result))
        texts = [p.text for p in doc.paragraphs]
        assert not any("titulo:" in t for t in texts)

    def test_bullet_list(self):
        md = "- Item 1\n- Item 2\n- Item 3"
        result = markdown_to_docx(md)
        doc = Document(io.BytesIO(result))
        bullets = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
        assert len(bullets) == 3
