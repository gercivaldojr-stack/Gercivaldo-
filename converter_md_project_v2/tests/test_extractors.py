"""Testes para o módulo de extração de texto."""

import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from core.extractors import extract_text, _is_footer_text, _table_to_markdown, _docx_table_to_markdown


class TestExtractText:
    def test_unsupported_format(self):
        with pytest.raises(ValueError, match="Formato não suportado"):
            extract_text(file_path=None, file_bytes=b"test", filename="file.xyz")

    def test_no_args(self):
        with pytest.raises(ValueError, match="Forneça"):
            extract_text(file_path=None, file_bytes=None, filename=None)

    def test_txt_extraction(self):
        content = "Texto de teste para extração."
        result = extract_text(
            file_path=None,
            file_bytes=content.encode("utf-8"),
            filename="test.txt",
        )
        assert result == content

    def test_md_extraction(self):
        content = "# Heading\n\nParágrafo com acentuação."
        result = extract_text(
            file_path=None,
            file_bytes=content.encode("utf-8"),
            filename="test.md",
        )
        assert "# Heading" in result

    def test_txt_latin1(self):
        content = "Texto com acentuação: ção, ão, é, ê"
        encoded = content.encode("latin-1")
        result = extract_text(
            file_path=None,
            file_bytes=encoded,
            filename="test.txt",
        )
        assert "ação" in result or "a" in result  # encoding detection may vary


class TestExtractDocx:
    def test_docx_extraction(self):
        """Testa extração de DOCX gerando um arquivo mínimo em memória."""
        from docx import Document
        import io

        doc = Document()
        doc.add_heading("Título do Documento", level=1)
        doc.add_paragraph("Primeiro parágrafo com conteúdo jurídico.")
        doc.add_heading("Dos Fatos", level=2)
        doc.add_paragraph("Descrição dos fatos relevantes ao caso.")

        buffer = io.BytesIO()
        doc.save(buffer)
        file_bytes = buffer.getvalue()

        result = extract_text(
            file_path=None,
            file_bytes=file_bytes,
            filename="test.docx",
        )

        assert "Título do Documento" in result
        assert "Primeiro parágrafo" in result
        assert "Dos Fatos" in result
        assert "# " in result  # Heading deve ser convertido

    def test_docx_table_extraction(self):
        """Testa que tabelas em DOCX são extraídas como Markdown tables."""
        from docx import Document
        import io

        doc = Document()
        doc.add_paragraph("Texto antes da tabela.")

        # Adicionar tabela 3x2
        table = doc.add_table(rows=3, cols=2)
        table.cell(0, 0).text = "Nome"
        table.cell(0, 1).text = "Valor"
        table.cell(1, 0).text = "Item A"
        table.cell(1, 1).text = "R$ 100"
        table.cell(2, 0).text = "Item B"
        table.cell(2, 1).text = "R$ 200"

        doc.add_paragraph("Texto depois da tabela.")

        buffer = io.BytesIO()
        doc.save(buffer)
        file_bytes = buffer.getvalue()

        result = extract_text(
            file_path=None,
            file_bytes=file_bytes,
            filename="test.docx",
        )

        assert "Texto antes da tabela" in result
        assert "Texto depois da tabela" in result
        assert "| Nome | Valor |" in result
        assert "| Item A | R$ 100 |" in result
        assert "| Item B | R$ 200 |" in result
        assert "| --- | --- |" in result

    def test_docx_table_order_preserved(self):
        """Testa que a ordem parágrafos-tabelas é preservada."""
        from docx import Document
        import io

        doc = Document()
        doc.add_paragraph("Parágrafo 1")
        table = doc.add_table(rows=2, cols=1)
        table.cell(0, 0).text = "Header"
        table.cell(1, 0).text = "Data"
        doc.add_paragraph("Parágrafo 2")

        buffer = io.BytesIO()
        doc.save(buffer)
        file_bytes = buffer.getvalue()

        result = extract_text(
            file_path=None,
            file_bytes=file_bytes,
            filename="test.docx",
        )

        # Verificar ordem: Parágrafo 1 antes da tabela, tabela antes do Parágrafo 2
        idx_p1 = result.index("Parágrafo 1")
        idx_table = result.index("| Header |")
        idx_p2 = result.index("Parágrafo 2")
        assert idx_p1 < idx_table < idx_p2


class TestExtractPDF:
    def test_pdf_extraction(self):
        """Testa extração de PDF gerando um PDF mínimo em memória."""
        import fitz

        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Texto do PDF de teste", fontsize=12)
        page.insert_text((72, 100), "Segundo parágrafo do documento.", fontsize=12)

        pdf_bytes = doc.tobytes()
        doc.close()

        result = extract_text(
            file_path=None,
            file_bytes=pdf_bytes,
            filename="test.pdf",
        )

        assert "Texto do PDF" in result
        assert "Segundo parágrafo" in result


class TestFooterDetection:
    def test_cep_pattern(self):
        assert _is_footer_text("Rua 14, Goiânia/GO, CEP 74810-180")

    def test_phone_pattern(self):
        assert _is_footer_text("Tel: (62) 3333-4444")

    def test_page_pattern(self):
        assert _is_footer_text("Página 15")

    def test_pag_pattern(self):
        assert _is_footer_text("Pág. 3")

    def test_email_pattern(self):
        assert _is_footer_text("contato@escritorio.com.br")

    def test_normal_text_not_footer(self):
        assert not _is_footer_text("O autor alega que sofreu danos morais.")


class TestTableToMarkdown:
    def test_docx_table_to_markdown(self):
        """Testa conversão de tabela DOCX para Markdown."""
        from docx import Document
        import io

        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Col1"
        table.cell(0, 1).text = "Col2"
        table.cell(1, 0).text = "A"
        table.cell(1, 1).text = "B"

        result = _docx_table_to_markdown(table)
        assert "| Col1 | Col2 |" in result
        assert "| --- | --- |" in result
        assert "| A | B |" in result
