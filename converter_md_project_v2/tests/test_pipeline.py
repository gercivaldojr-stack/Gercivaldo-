"""Testes para o pipeline principal."""

import io
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from core.pipeline import convert_batch, convert_document


class TestConvertDocument:
    def test_txt_conversion(self):
        content = "DOS FATOS\n\nO autor alega que...\n\nDOS PEDIDOS\n\nRequer-se a condenação."
        result = convert_document(
            file_bytes=content.encode("utf-8"),
            filename="peticao.txt",
            mode="forense",
        )
        assert result.success
        assert "## DOS FATOS" in result.markdown
        assert "## DOS PEDIDOS" in result.markdown

    def test_txt_doutrina(self):
        content = "CAPÍTULO I - Introdução\n\nTexto introdutório.\n\n1.1 Conceitos\n\nDefinições."
        result = convert_document(
            file_bytes=content.encode("utf-8"),
            filename="doutrina.txt",
            mode="doutrina",
        )
        assert result.success
        assert "# CAPÍTULO I" in result.markdown
        assert "### 1.1" in result.markdown

    def test_empty_document(self):
        result = convert_document(
            file_bytes=b"",
            filename="vazio.txt",
        )
        assert not result.success
        assert "vazio" in result.error.lower() or "empty" in result.error.lower()

    def test_piece_separation(self):
        content = (
            "PETIÇÃO INICIAL\n\nConteúdo da petição.\n\n"
            "CONTESTAÇÃO\n\nConteúdo da contestação."
        )
        result = convert_document(
            file_bytes=content.encode("utf-8"),
            filename="autos.txt",
            mode="forense",
            separate=True,
        )
        assert result.success
        # 3 peças: frontmatter (Documento Principal) + PETIÇÃO INICIAL + CONTESTAÇÃO
        assert len(result.pieces) >= 2
        titles = [p["title"] for p in result.pieces]
        assert any("PETIÇÃO" in t for t in titles)
        assert any("CONTESTAÇÃO" in t for t in titles)

    def test_docx_conversion(self):
        from docx import Document

        doc = Document()
        doc.add_heading("Sentença", level=1)
        doc.add_paragraph("O juiz decide...")

        buffer = io.BytesIO()
        doc.save(buffer)

        result = convert_document(
            file_bytes=buffer.getvalue(),
            filename="sentenca.docx",
            mode="forense",
        )
        assert result.success
        assert "Sentença" in result.markdown

    def test_stats_populated(self):
        content = "Texto simples para teste de estatísticas."
        result = convert_document(
            file_bytes=content.encode("utf-8"),
            filename="stats.txt",
        )
        assert result.stats.get("chars_raw", 0) > 0
        assert result.stats.get("chars_final", 0) > 0


class TestConvertBatch:
    def test_batch_conversion(self):
        files = [
            {"file_bytes": b"DOS FATOS\n\nTexto 1.", "filename": "doc1.txt"},
            {"file_bytes": b"DOS PEDIDOS\n\nTexto 2.", "filename": "doc2.txt"},
        ]
        results = convert_batch(files, mode="forense")
        assert len(results) == 2
        assert all(r.success for r in results)

    def test_batch_with_error(self):
        files = [
            {"file_bytes": b"Texto ok.", "filename": "ok.txt"},
            {"file_bytes": b"data", "filename": "bad.xyz"},
        ]
        results = convert_batch(files, mode="forense")
        assert len(results) == 2
        assert results[0].success
        assert not results[1].success
